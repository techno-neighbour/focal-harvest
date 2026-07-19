import os
import re
import utils

logger = utils.setup_logging()

def _clean_markdown_text(text: str) -> str:
    """Strips Markdown links, bold/italic markers for clean plain text rendering."""
    # Convert [text](url) to text (url)
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', text)
    # Strip bold / italic stars
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text.strip()

def _flush_pdf_table(table_rows, story, body_style, ParagraphStyle, Paragraph, Table, TableStyle, Spacer, colors):
    """Renders a parsed Markdown table into a styled ReportLab Table object."""
    if not table_rows:
        return
    formatted_table_data = []
    for r_idx, row in enumerate(table_rows):
        formatted_row = []
        for cell in row:
            clean_c = _clean_markdown_text(cell)
            cell_style = ParagraphStyle(
                f'Cell_{r_idx}_{len(formatted_row)}',
                parent=body_style,
                fontSize=8.5,
                leading=11,
                fontName='Helvetica-Bold' if r_idx == 0 else 'Helvetica'
            )
            formatted_row.append(Paragraph(clean_c, cell_style))
        formatted_table_data.append(formatted_row)
    
    if formatted_table_data:
        cols_count = len(table_rows[0])
        if cols_count == 2:
            col_widths = [200, 340]
        elif cols_count == 3:
            col_widths = [50, 220, 270]
        elif cols_count == 4:
            col_widths = [30, 180, 270, 60]
        else:
            col_widths = [540.0 / cols_count] * cols_count
            
        t = Table(formatted_table_data, colWidths=col_widths)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#E0F2FE')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor('#0F172A')),
            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
            ('TOPPADDING', (0,0), (-1,-1), 4),
            ('BOTTOMPADDING', (0,0), (-1,-1), 4),
        ]))
        story.append(t)
        story.append(Spacer(1, 10))

def export_markdown_to_pdf(markdown_content: str, output_filepath: str) -> bool:
    """
    Exports a Markdown report string to a clean, styled PDF document using ReportLab.
    Returns True if generation succeeds, False if reportlab is missing or fails.
    """
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib import colors
    except ImportError:
        logger.warning("reportlab package not installed. Skipping PDF export. Run 'pip install reportlab' to enable PDF exports.")
        return False

    try:
        os.makedirs(os.path.dirname(output_filepath), exist_ok=True)
        doc = SimpleDocTemplate(
            output_filepath,
            pagesize=letter,
            rightMargin=36,
            leftMargin=36,
            topMargin=36,
            bottomMargin=36
        )

        styles = getSampleStyleSheet()
        
        # Custom Paragraph Styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            leading=22,
            textColor=colors.HexColor('#1E3A8A'),
            spaceAfter=12,
            fontName='Helvetica-Bold'
        )
        h2_style = ParagraphStyle(
            'CustomH2',
            parent=styles['Heading2'],
            fontSize=14,
            leading=18,
            textColor=colors.HexColor('#1E3A8A'),
            spaceBefore=14,
            spaceAfter=8,
            fontName='Helvetica-Bold'
        )
        h3_style = ParagraphStyle(
            'CustomH3',
            parent=styles['Heading3'],
            fontSize=11,
            leading=15,
            textColor=colors.HexColor('#2563EB'),
            spaceBefore=10,
            spaceAfter=6,
            fontName='Helvetica-Bold'
        )
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=9.5,
            leading=13.5,
            textColor=colors.HexColor('#1F2937'),
            spaceAfter=6,
            fontName='Helvetica'
        )
        bullet_style = ParagraphStyle(
            'CustomBullet',
            parent=body_style,
            leftIndent=12,
            spaceAfter=4
        )

        story = []
        lines = markdown_content.splitlines()
        in_table = False
        table_rows = []

        for line in lines:
            line_str = line.strip()

            # Handle Markdown Tables
            if line_str.startswith("|") and line_str.endswith("|"):
                in_table = True
                # Check for table header separator line
                if re.match(r'^\|[\s:-|-]+\|$', line_str):
                    continue
                cells = [c.strip() for c in line_str.split("|")[1:-1]]
                table_rows.append(cells)
                continue
            elif in_table:
                # Flush table when non-table line is encountered
                _flush_pdf_table(table_rows, story, body_style, ParagraphStyle, Paragraph, Table, TableStyle, Spacer, colors)
                in_table = False
                table_rows = []

            if not line_str:
                story.append(Spacer(1, 4))
                continue

            if line_str.startswith("# "):
                clean_text = _clean_markdown_text(line_str[2:])
                story.append(Paragraph(clean_text, title_style))
            elif line_str.startswith("## "):
                clean_text = _clean_markdown_text(line_str[3:])
                story.append(Paragraph(clean_text, h2_style))
            elif line_str.startswith("### "):
                clean_text = _clean_markdown_text(line_str[4:])
                story.append(Paragraph(clean_text, h3_style))
            elif line_str.startswith("- ") or line_str.startswith("* "):
                clean_text = _clean_markdown_text(line_str[2:])
                story.append(Paragraph(f"• {clean_text}", bullet_style))
            else:
                clean_text = _clean_markdown_text(line_str)
                story.append(Paragraph(clean_text, body_style))

        # Flush table if document ends on a table
        if in_table and table_rows:
            _flush_pdf_table(table_rows, story, body_style, ParagraphStyle, Paragraph, Table, TableStyle, Spacer, colors)

        doc.build(story)
        logger.info("Successfully generated PDF report at '%s'", output_filepath)
        return True
    except Exception as e:
        logger.error("Failed to generate PDF report at '%s': %s", output_filepath, str(e))
        return False

def _flush_docx_table(table_rows, doc, RGBColor):
    """Renders a parsed Markdown table into a python-docx Table object."""
    if not table_rows:
        return
    cols_count = len(table_rows[0])
    t = doc.add_table(rows=len(table_rows), cols=cols_count)
    t.style = 'Table Grid'
    for r_idx, row in enumerate(table_rows):
        for c_idx, cell in enumerate(row[:cols_count]):
            clean_c = _clean_markdown_text(cell)
            cell_para = t.cell(r_idx, c_idx).paragraphs[0]
            cell_para.text = clean_c
            if r_idx == 0:
                for run in cell_para.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(30, 58, 138)
    doc.add_paragraph("") # Spacer after table

def export_markdown_to_docx(markdown_content: str, output_filepath: str) -> bool:
    """
    Exports a Markdown report string to a clean, styled Word (.docx) document using python-docx.
    Returns True if generation succeeds, False if python-docx is missing or fails.
    """
    try:
        import docx
        from docx.shared import Inches, Pt, RGBColor
    except ImportError:
        logger.warning("python-docx package not installed. Skipping DOCX export. Run 'pip install python-docx' to enable Word exports.")
        return False

    try:
        os.makedirs(os.path.dirname(output_filepath), exist_ok=True)
        doc = docx.Document()

        # Page Setup
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        lines = markdown_content.splitlines()
        in_table = False
        table_rows = []

        for line in lines:
            line_str = line.strip()

            # Handle Markdown Tables
            if line_str.startswith("|") and line_str.endswith("|"):
                in_table = True
                if re.match(r'^\|[\s:-|-]+\|$', line_str):
                    continue
                cells = [c.strip() for c in line_str.split("|")[1:-1]]
                table_rows.append(cells)
                continue
            elif in_table:
                # Flush table when non-table line is encountered
                _flush_docx_table(table_rows, doc, RGBColor)
                in_table = False
                table_rows = []

            if not line_str:
                continue

            if line_str.startswith("# "):
                clean_text = _clean_markdown_text(line_str[2:])
                h = doc.add_heading(clean_text, level=1)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(30, 58, 138)
            elif line_str.startswith("## "):
                clean_text = _clean_markdown_text(line_str[3:])
                h = doc.add_heading(clean_text, level=2)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(30, 58, 138)
            elif line_str.startswith("### "):
                clean_text = _clean_markdown_text(line_str[4:])
                h = doc.add_heading(clean_text, level=3)
                for run in h.runs:
                    run.font.color.rgb = RGBColor(37, 99, 235)
            elif line_str.startswith("- ") or line_str.startswith("* "):
                clean_text = _clean_markdown_text(line_str[2:])
                doc.add_paragraph(clean_text, style='List Bullet')
            else:
                clean_text = _clean_markdown_text(line_str)
                doc.add_paragraph(clean_text)

        # Flush table if document ends on a table
        if in_table and table_rows:
            _flush_docx_table(table_rows, doc, RGBColor)

        doc.save(output_filepath)
        logger.info("Successfully generated DOCX report at '%s'", output_filepath)
        return True
    except Exception as e:
        logger.error("Failed to generate DOCX report at '%s': %s", output_filepath, str(e))
        return False
